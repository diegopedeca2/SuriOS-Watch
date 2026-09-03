from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\WristOS")
OUTPUT = ROOT / "output" / "SPRINT_030_APK" / "PIP-SuriOS_ALPHA_TEST_GUIDE_SPRINT_030.docx"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = "666666"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, name="Calibri", size=11, color="000000", bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("PIP-SuriOS  ·  GUÍA ALPHA TESTER  ·  SPRINT 030")
    set_font(run, size=8, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Documento de pruebas · Versión de distribución 1.0")
    set_font(run, size=8, color=MUTED)


def add_title(doc: Document, text: str, subtitle: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_font(run, size=26, color=DARK_BLUE, bold=True)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(subtitle)
    set_font(run, size=12, color=MUTED, italic=True)


def add_body(doc, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix), bold=True)
        set_font(p.add_run(text[len(bold_prefix):]))
    else:
        set_font(p.add_run(text))
    return p


def add_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    set_font(p.add_run(text))
    return p


def add_number(doc, text: str):
    p = doc.add_paragraph(style="List Number")
    set_font(p.add_run(text))
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        cell.text = ""
        p = cell.paragraphs[0]
        set_font(p.add_run(header), size=10, color=DARK_BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = ""
            p = cells[index].paragraphs[0]
            set_font(p.add_run(value), size=9.5)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_test_block(doc, title, purpose, steps, expected):
    doc.add_heading(title, level=2)
    add_body(doc, purpose, bold_prefix="Objetivo: ")
    for step in steps:
        add_number(doc, step)
    add_body(doc, "Resultado esperado: " + expected, bold_prefix="Resultado esperado: ")


def main() -> None:
    doc = Document()
    style_document(doc)
    add_title(
        doc,
        "PIP-SuriOS · Guía de pruebas Alpha",
        "Sprint 030 · Documento para FENRIR, ALTAMIRA y CHECHU · 03/09/2026",
    )
    add_body(doc, "Gracias por ser las primeras personas en probar esta versión y por ayudarnos a detectar problemas reales de uso. No hace falta saber programación: basta con usar la aplicación con normalidad y anotar lo que ocurra.")

    doc.add_heading("1. Qué versión debe usar cada tester", level=1)
    add_table(
        doc,
        ["Tester / APK", "Mapas incluidos", "Observación"],
        [
            ("FENRIR v3.0", "NAVY7 y TESTING", "TESTING conserva el campo específico de FENRIR."),
            ("ALTAMIRA v3.0", "NAVY7 y TESTING", "TESTING está centrado en 40.34897942140349, -3.818235386395919."),
            ("CHECHU v3.0", "NAVY7 y TESTING", "TESTING está centrado en 40.433753, -3.625904."),
            ("MAIN v3.0", "HOME, NAVY7 y OFFICE", "Es la APK principal de trabajo del A56; no es la APK Alpha asignada."),
        ],
        [2050, 1850, 5460],
    )
    add_body(doc, "Las cuatro aplicaciones pueden estar instaladas a la vez porque cada una tiene un identificador independiente. En Android, comprueba que abres la aplicación con el nombre de tu tester y el icono correspondiente: PIP-F, PIP-A o PIP-C.")
    add_body(doc, "Importante: estas son versiones de prueba. No deben utilizarse para tomar decisiones reales de seguridad, navegación o exposición a radiación.")

    doc.add_heading("2. Ficha del dispositivo y de la prueba", level=1)
    add_table(
        doc,
        ["Dato", "Anotar antes de empezar"],
        [
            ("Tester", "__________________________________________________"),
            ("Modelo de teléfono", "__________________________________________________"),
            ("Versión de Android", "__________________________________________________"),
            ("APK utilizada", "__________________________________________________"),
            ("Fecha y lugar", "__________________________________________________"),
            ("¿Se usó reloj / PROBE?", "Sí / No · Modelo: ______________________________"),
        ],
        [2700, 6660],
    )

    doc.add_heading("3. Preparación y permisos", level=1)
    add_body(doc, "Al abrir la aplicación por primera vez, acepta solo los permisos que quieras probar. Si rechazas uno, anótalo: también es información útil.")
    add_table(
        doc,
        ["Permiso / acceso", "Para qué lo usa la aplicación"],
        [
            ("Bluetooth: dispositivos cercanos", "P.R.S., conexión con el Watch 2 / PROBE y lectura de dispositivos Bluetooth."),
            ("Ubicación precisa o aproximada", "GPS del teléfono, posición en mapas y parte del escaneo Bluetooth de Android."),
            ("Cámara", "Morse: se utiliza el flash trasero para transmitir señales luminosas."),
            ("Google Maps / CivTAK", "Son aplicaciones externas opcionales para MAP OPERATION; si no están instaladas, anotar el mensaje."),
        ],
        [2700, 6660],
    )
    add_bullet(doc, "Activa el volumen multimedia antes de probar RADS.")
    add_bullet(doc, "Para las pruebas de Bluetooth, activa Bluetooth y ubicación del teléfono.")
    add_bullet(doc, "Para las pruebas de mapa, concede ubicación solo si quieres comprobar el GPS; NAVY7 y TESTING deben poder abrirse sin conexión de datos.")

    doc.add_heading("4. Pruebas obligatorias", level=1)
    add_test_block(doc, "4.1 Inicio, identificación y carga", "Comprobar que el arranque se entiende y que el texto permanece visible el tiempo suficiente.", [
        "Abrir la APK asignada desde el icono del tester.",
        "Completar el lector / identificación y observar LOG-IN CONFIRM, WELCOME y la pantalla de carga.",
        "Entrar en HOMESCREEN y volver a abrir la aplicación una segunda vez.",
    ], "No hay cierres, el texto se puede leer y la aplicación llega a HOMESCREEN.")

    add_test_block(doc, "4.2 SET-UP del operador y equipamiento", "Comprobar que las casillas manuales guardan el equipo particular.", [
        "Entrar en SET-UP OPERATOR y guardar un ID. Confirmar que no aparece CALLSIGN.",
        "En PRIMARY WEAPON y SECONDARY WEAPON escribir dos réplicas diferentes. Pulsar APPLY después de cada una.",
        "Repetir con ACCESORIES, FRONT PANEL y UNIFORM. Cada APPLY debe guardar el texto y dejar la casilla vacía para introducir otro elemento.",
        "En HEADGEAR introducir primero el nombre y después varios componentes usando sus casillas y APPLY.",
        "Salir y volver a entrar en SET-UP para comprobar que los datos siguen guardados.",
    ], "Los textos no se mezclan, APPLY limpia la casilla y los datos permanecen al navegar.")

    add_test_block(doc, "4.3 DATA y CURRENT GEAR", "Comprobar que el equipo guardado se muestra con nombres concretos.", [
        "Abrir DATA y revisar las categorías de armas, accesorios, frontal, uniforme y headgear.",
        "Editar o borrar una opción y comprobar el cambio.",
        "Abrir CURRENT GEAR y confirmar que muestra el equipo activo, sin opciones genéricas como OPTION 1.",
    ], "DATA y CURRENT GEAR reflejan el contenido introducido y no pierden elementos al volver atrás.")

    add_test_block(doc, "4.4 RADS", "Comprobar el medidor y valorar el nuevo sonido de ráfaga/crujido.", [
        "Abrir TOOLS > RADS y probar los niveles LOW, HIGH y CRITICAL.",
        "Escuchar si el sonido parece una ráfaga irregular de microdescargas, no una fila de clics iguales.",
        "Bajar el nivel a 0 y volver a subirlo sin salir de la herramienta.",
        "Repetir la prueba varias veces y anotar si el ritmo se repite demasiado o si hay cortes.",
    ], "El sonido se inicia al detectar nivel, se detiene en 0 y vuelve automáticamente al subir. El ritmo cambia de forma natural.")

    add_test_block(doc, "4.5 MAP > TERRAIN", "Comprobar los mapas específicos de cada APK.", [
        "Abrir TOOLS > MAP > TERRAIN y revisar la lista de mapas.",
        "Confirmar que HOME y OFFICE no aparecen en las APK Alpha.",
        "En FENRIR abrir TESTING y comprobar que el centro corresponde a la zona indicada.",
        "En ALTAMIRA abrir TESTING y comprobar que el centro corresponde a 40.34897942140349, -3.818235386395919.",
        "En CHECHU abrir TESTING y comprobar que el centro corresponde a 40.433753, -3.625904.",
        "Probar pan, zoom, brújula, GPS y la creación de un punto de respawn o zona RAD solo si procede.",
    ], "La lista coincide con la tabla de esta guía, los mapas no se cierran y el comportamiento táctil es comprensible.")

    add_test_block(doc, "4.6 P.R.S. y Bluetooth", "Comprobar el escaneo, el seguimiento, las exclusiones y los estados de conexión.", [
        "Abrir PROXIMITY RADIO SCANNER y revisar SENTRY, TRACKER, DEVICES y USER GUIDE.",
        "En SENTRY probar PIP y PIP + PROBE. Anotar qué ocurre sin dispositivos cercanos y qué datos aparecen cuando hay un objetivo controlado.",
        "En TRACKER elegir ONLY PIP-BOY o PIP-BOY + PROBE, seleccionar primero el mapa TERRAIN y después el objetivo detectado.",
        "Comprobar que al entrar en la pantalla del objetivo la lectura empieza sola. No hay que buscar un botón START.",
        "Dejar la pantalla abierta y comprobar que RAW puede cambiar de inmediato, mientras SMOOTH, TREND, SAMPLES y CONFIDENCE se completan con el tiempo.",
        "Esperar al menos 15 segundos antes de valorar la tendencia. Usar BACK para salir y comprobar que la sesión termina.",
        "En DEVICES añadir, editar y borrar un dispositivo a omitir. Comprobar que SENTRY y TRACKER lo respetan.",
        "En ONLY PIP-BOY no hace falta vincular otro dispositivo. Para PIP-BOY + PROBE, comprobar que el Watch 2 está emparejado y conectado.",
    ], "Los menús cargan, el seguimiento comienza y termina de forma entendible, y los fallos de permisos o conexión muestran un estado útil.")

    add_test_block(doc, "4.7 P.R.S. - prueba P01: señal estable por distancia", "Obtener una referencia sencilla entre señal y distancia real en un entorno controlado. La distancia se mide fuera de la aplicación; P.R.S. no calcula metros.", [
        "Usar un dispositivo Bluetooth controlado como objetivo y mantenerlo quieto. Mantener también el A56 quieto durante cada tramo.",
        "Probar, si es seguro y posible, aproximadamente 1 m, 3 m, 5 m y 10 m. Si el lugar no permite esas distancias, anotar las distancias disponibles.",
        "En cada distancia, mantener TRACKER abierto unos 30 segundos. Esperar al menos 15 segundos antes de juzgar TREND o CONFIDENCE.",
        "Cada aproximadamente 3 segundos, anotar una fila en el CSV: distancia real, RAW, SMOOTH, TREND, banda, SAMPLES, CONFIDENCE y estado de BLE/GPS.",
        "Repetir una distancia que parezca inestable para comprobar si el resultado cambia.",
    ], "El CSV contiene varias filas por distancia y permite comparar la señal sin afirmar que exista una conversión fiable a metros.")

    add_test_block(doc, "4.8 P.R.S. - prueba P02: acercamiento y alejamiento", "Comprobar si la tendencia cambia de forma coherente cuando el receptor se acerca o se aleja del objetivo.", [
        "Colocarse a una distancia inicial conocida y mantener TRACKER abierto hasta que haya historial suficiente.",
        "Caminar lentamente hacia el objetivo durante al menos 30 segundos y anotar las filas como APPROACH.",
        "Volver caminando lentamente y alejarse durante al menos 30 segundos; anotar las filas como AWAY.",
        "Anotar también el movimiento del receptor, el rumbo aproximado y cualquier obstáculo o persona entre los dispositivos.",
    ], "La tendencia puede mostrar APPROACHING o MOVING AWAY, pero se acepta que haya ruido, retraso o WAITING mientras faltan muestras.")

    add_test_block(doc, "4.9 P.R.S. - prueba P03: orientación y obstáculos", "Separar los cambios producidos por la distancia de los cambios producidos por el cuerpo, la orientación o los obstáculos.", [
        "Elegir una distancia fija y repetir la lectura con el teléfono en la mano, en el bolsillo y sujeto en otra orientación.",
        "Repetir con línea de vista libre y después con un obstáculo seguro, como una pared o una persona, sin poner a nadie en riesgo.",
        "Mantener cada condición unos 30 segundos y escribir la condición en environment u obstacles del CSV.",
    ], "Los resultados muestran variación real del RSSI y dejan una nota suficiente para saber qué condición produjo cada cambio.")

    add_test_block(doc, "4.10 P.R.S. - prueba P04: desaparición y caducidad", "Comprobar qué ocurre cuando el objetivo deja de emitir o queda fuera del alcance.", [
        "Con el objetivo visible, anotar una fila normal y después apagarlo o alejarlo de forma controlada.",
        "Mantener TRACKER abierto al menos 20 segundos y anotar cuándo deja de aparecer o cambia de estado.",
        "Volver a encender o acercar el objetivo y comprobar si vuelve a aparecer sin cerrar la pantalla.",
    ], "El contacto puede tardar en desaparecer; la configuración actual considera caducado un contacto sin observaciones recientes después de aproximadamente 15 segundos.")

    add_test_block(doc, "4.11 P.R.S. - prueba P05: PROBE", "Comprobar el aporte del Watch 2 sin confundir la posición del receptor con la del objetivo.", [
        "Emparejar y conectar el Watch 2 antes de abrir PIP-BOY + PROBE.",
        "Repetir una distancia corta y una distancia larga de P01, anotando source como A56 o WATCH2 según la lectura.",
        "Anotar cualquier reconexión, pérdida de batería o cambio de estado en probe_status y notes.",
        "No interpretar la posición mostrada del Watch 2 como coordenada del objetivo: es la posición del receptor que aporta la lectura.",
    ], "Las lecturas indican su fuente y el estado de PROBE es entendible, incluso si el reloj no puede conectarse o se desconecta.")

    doc.add_heading("5. Registro estructurado de datos P.R.S.", level=1)
    add_body(doc, "Rellena el archivo PRS_FIELD_DATA_TEMPLATE.csv que acompaña a esta guía. Está separado por punto y coma (;) para que sea más fácil abrirlo con Excel en configuración española. Cada fila representa una observación aproximadamente cada 3 segundos, no una nueva sesión.")
    add_body(doc, "Los campos más importantes para esta primera campaña son:")
    for text in [
        "test_id: usa P01, P02, P03, P04 o P05 según la prueba.",
        "mode y source: ONLY_PIP_BOY o PIP_BOY_PROBE; A56 o WATCH2.",
        "actual_distance_m: distancia aproximada medida fuera de la aplicación. Déjalo vacío si no se puede estimar.",
        "raw_rssi_dbm, smoothed_rssi_dbm, trend, proximity_band, samples y confidence: copia los valores que aparecen en pantalla.",
        "environment, obstacles, receiver_movement y notes: describe las condiciones que pueden explicar cambios.",
        "target_identifier_optional: puede quedar vacío o usar un alias local. No hace falta compartir una dirección MAC.",
    ]:
        add_bullet(doc, text)
    add_body(doc, "Entrega el CSV junto con capturas o vídeo cuando haya un comportamiento extraño. Indica siempre la APK, el modelo del teléfono, el test_id y el momento aproximado de la incidencia.")

    doc.add_heading("6. Pruebas adicionales rápidas", level=1)
    for text in [
        "STATUS: revisar que el estado del operador y del equipamiento sea legible.",
        "INVENTORY y STORAGE: crear, consultar y borrar un elemento de prueba si el flujo lo permite.",
        "COMMS / MORSE: transmitir un texto corto con flash y comprobar que el control vuelve a su estado normal.",
        "MAP OPERATION: probar el acceso a Google Maps o CivTAK y anotar si la aplicación externa no está instalada.",
        "Rotación, toque rápido, volver atrás y reabrir: anotar cualquier pantalla en blanco, desplazamiento inesperado o pérdida de datos.",
    ]:
        add_bullet(doc, text)

    doc.add_heading("7. Riesgos y límites conocidos", level=1)
    for text in [
        "Los mapas offline solo cubren las zonas incluidas en cada APK. TESTING de ALTAMIRA y CHECHU son campos específicos de sus APK.",
        "La posición GPS, brújula y Bluetooth dependen del teléfono, permisos, cobertura, batería y entorno físico.",
        "La conexión con Watch 2 / PROBE no se puede validar completamente sin disponer del hardware compatible.",
        "El sonido de RADS depende del volumen multimedia y del altavoz o auriculares del teléfono; la valoración de realismo es subjetiva.",
        "Google Maps y CivTAK son aplicaciones externas: su ausencia o cambios propios pueden afectar MAP OPERATION.",
        "La aplicación está en fase Alpha: puede haber textos provisionales, cambios visuales, pérdida de datos de prueba o cierres inesperados.",
        "RADS, P.R.S. y los mapas son herramientas de simulación/prueba; no sustituyen instrumentos profesionales ni procedimientos reales.",
    ]:
        add_bullet(doc, text)

    doc.add_heading("8. Cómo enviar el feedback", level=1)
    add_body(doc, "Para cada problema, envía una descripción corta con estos datos. Una grabación de pantalla, captura o vídeo del sonido ayuda mucho.")
    add_table(
        doc,
        ["Campo", "Respuesta"],
        [
            ("Pantalla / función", "__________________________________________________"),
            ("Pasos para repetirlo", "1) __________________  2) __________________  3) __________________"),
            ("Qué esperaba", "__________________________________________________"),
            ("Qué ocurrió", "__________________________________________________"),
            ("¿Se repite?", "Siempre / A veces / Una vez"),
            ("Gravedad", "Bloquea / Molesto / Menor / Sugerencia"),
            ("Captura, vídeo o audio", "Adjunto / No disponible"),
        ],
        [2700, 6660],
    )
    add_body(doc, "En RADS, añade además: nivel aproximado, si el sonido estaba activo, volumen multimedia y si volvió a sonar después de bajar a 0.")
    add_body(doc, "En mapas, añade: nombre del mapa, si había conexión de datos, si se había concedido ubicación y si el fallo afectó a pan, zoom, brújula, GPS u overlays.")

    doc.add_heading("9. Criterio de cierre de la prueba", level=1)
    add_body(doc, "La prueba se considera completa cuando se han recorrido las secciones 4.1 a 4.11, se ha rellenado el CSV de P.R.S. si se dispone de un objetivo controlado, se ha anotado el modelo del dispositivo y se ha enviado feedback incluso si todo funciona. Los comentarios positivos también son útiles: indican qué partes no debemos romper en la siguiente versión.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
